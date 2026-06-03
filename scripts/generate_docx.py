#!/usr/bin/env python3
"""
河南科技大学本科毕业设计（论文）Word 文档生成辅助脚本

基于河南科技大学论文格式规范，使用 python-docx 生成符合要求的 .docx 论文框架。
本脚本生成包含所有必要部分的论文模板，用户可在此基础上填充内容。

依赖：pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import sys
import json


def set_font(run, cn_font="宋体", en_font="Times New Roman", size=None, bold=False, color=None):
    """设置 run 的中英文字体"""
    run.font.name = en_font
    run.font.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 设置中文字体
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{cn_font}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), cn_font)


def set_paragraph_format(paragraph, first_line_indent=None, line_spacing=1.5,
                         space_before=0, space_after=0, alignment=None):
    """设置段落格式"""
    pf = paragraph.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Pt(first_line_indent)
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if alignment is not None:
        pf.alignment = alignment


def add_page_break(doc):
    """添加分页符"""
    doc.add_page_break()


def create_cover_page(doc, info):
    """创建封面页"""
    # 添加空段用于调整位置
    for _ in range(3):
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 英文校名
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("HENAN UNIVERSITY OF SCIENCE & TECHNOLOGY")
    set_font(run, cn_font="Times New Roman", en_font="Times New Roman", size=16, bold=True)

    # 空行
    doc.add_paragraph()

    # 毕业设计（论文）标题
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("毕 业 设 计（论 文）")
    set_font(run, cn_font="黑体", size=26, bold=True)

    # 空行
    for _ in range(3):
        doc.add_paragraph()

    # 信息字段
    fields = [
        ("题    目", info.get("title", "")),
        ("姓    名", info.get("name", "")),
        ("学    院", info.get("college", "")),
        ("专    业", info.get("major", "")),
        ("指导教师", info.get("advisor", "")),
    ]

    for label, value in fields:
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(f"{label}    {value}")
        set_font(run, cn_font="宋体", size=14)

    # 空行
    for _ in range(3):
        doc.add_paragraph()

    # 日期
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    date_str = info.get("date", "2025 年 5 月 25 日")
    run = p.add_run(date_str)
    set_font(run, cn_font="宋体", size=14)


def create_declaration_page(doc, info):
    """创建学位论文写作声明页"""
    # 声明标题
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
    run = p.add_run("学位论文写作声明")
    set_font(run, cn_font="黑体", size=16, bold=True)

    # 声明正文
    declaration_text = (
        "本人郑重声明：所呈交的学位论文，是本人在导师的指导下，独立进行研究"
        "工作所取得的成果。除文中已经注明引用的内容外，本论文不含任何其他个人或"
        "集体已经发表或撰写过的作品或成果。对本文的研究做出重要贡献的个人和集体，"
        "均已在文中以明确方式标明。本声明的法律结果由本人承担。"
    )
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=24, line_spacing=1.5)
    run = p.add_run(declaration_text)
    set_font(run, cn_font="宋体", size=12)

    # 签名行
    doc.add_paragraph()
    date_str = info.get("date", "2025年5月25日")
    p = doc.add_paragraph()
    run = p.add_run(f"论文作者签名：            日期：{date_str}")
    set_font(run, cn_font="宋体", size=12)

    # 空行
    for _ in range(2):
        doc.add_paragraph()

    # 授权说明标题
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
    run = p.add_run("学位论文使用授权说明")
    set_font(run, cn_font="黑体", size=16, bold=True)

    # 授权正文
    auth_text = (
        "本人完全了解河南科技大学关于收集、保存、使用学位论文的规定，即：按"
        "照学校要求提交学位论文的印刷本和电子版本；学校有权保存学位论文的印刷本"
        "和电子版，并提供目录检索与阅览服务；学校可以采用影印、缩印、数字化或其"
        "它复制手段保存论文；在不以赢利为目的的前提下，学校可以将学位论文编入有"
        "关数据库,提供网上服务。（保密论文在解密后遵守此规定）"
    )
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=24, line_spacing=1.5)
    run = p.add_run(auth_text)
    set_font(run, cn_font="宋体", size=12)

    # 签名行
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"论文作者签名：              导师签名：")
    set_font(run, cn_font="宋体", size=12)
    p = doc.add_paragraph()
    run = p.add_run(f"日期：{date_str}")
    set_font(run, cn_font="宋体", size=12)


def create_chinese_abstract(doc, title, abstract_text, keywords):
    """创建中文摘要页"""
    # 页眉模拟（通过段落）
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    run = p.add_run("河南科技大学毕业设计说明书（论文）")
    set_font(run, cn_font="宋体", size=9)

    # 论文标题
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
    run = p.add_run(title)
    set_font(run, cn_font="黑体", size=16, bold=True)

    # 摘要标题
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
    run = p.add_run("摘 要")
    set_font(run, cn_font="黑体", size=16, bold=True)

    # 摘要正文
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=24, line_spacing=1.5)
    run = p.add_run(abstract_text)
    set_font(run, cn_font="宋体", size=12)

    # 关键词
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("关键词：")
    set_font(run, cn_font="黑体", size=12, bold=True)
    run = p.add_run(keywords)
    set_font(run, cn_font="宋体", size=12)


def create_english_abstract(doc, title_en, abstract_text_en, keywords_en):
    """创建英文摘要页"""
    # 页眉：英文标题
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    run = p.add_run(title_en.upper())
    set_font(run, en_font="Times New Roman", size=9)

    # 英文标题
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
    run = p.add_run(title_en.upper())
    set_font(run, en_font="Times New Roman", size=16, bold=True)

    # ABSTRACT
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
    run = p.add_run("ABSTRACT")
    set_font(run, en_font="Times New Roman", size=16, bold=True)

    # 英文摘要正文
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=24, line_spacing=1.5)
    run = p.add_run(abstract_text_en)
    set_font(run, en_font="Times New Roman", size=12)

    # KEY WORDS
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("KEY WORDS: ")
    set_font(run, en_font="Times New Roman", size=12, bold=True)
    run = p.add_run(keywords_en)
    set_font(run, en_font="Times New Roman", size=12)


def add_chapter_heading(doc, chapter_num, title):
    """添加章标题"""
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=12)
    run = p.add_run(f"第{chapter_num}章 {title}")
    set_font(run, cn_font="黑体", size=18, bold=True)


def add_section_heading(doc, section_num, title):
    """添加节标题"""
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=12, space_after=6)
    run = p.add_run(f"{section_num} {title}")
    set_font(run, cn_font="黑体", size=14, bold=True)


def add_subsection_heading(doc, subsection_num, title):
    """添加子节标题"""
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=6, space_after=6)
    run = p.add_run(f"{subsection_num} {title}")
    set_font(run, cn_font="黑体", size=12, bold=True)


def add_body_text(doc, text):
    """添加正文段落"""
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=24, line_spacing=1.5)
    run = p.add_run(text)
    set_font(run, cn_font="宋体", size=12)


def generate_thesis(info):
    """生成完整的论文框架"""
    doc = Document()

    # 设置默认页面格式
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    # 1. 封面
    create_cover_page(doc, info)
    add_page_break(doc)

    # 2. 空白页
    add_page_break(doc)

    # 3. 声明页
    create_declaration_page(doc, info)
    add_page_break(doc)

    # 4. 空白页
    add_page_break(doc)

    # 5. 中文摘要
    create_chinese_abstract(
        doc,
        info.get("title", ""),
        info.get("abstract_cn", "[在此填写中文摘要]"),
        info.get("keywords_cn", "[关键词1]；[关键词2]；[关键词3]")
    )
    add_page_break(doc)

    # 6. 英文摘要
    create_english_abstract(
        doc,
        info.get("title_en", ""),
        info.get("abstract_en", "[Fill in English abstract here]"),
        info.get("keywords_en", "[keyword1]; [keyword2]; [keyword3]")
    )
    add_page_break(doc)

    # 7. 目录（占位）
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=24)
    run = p.add_run("目 录")
    set_font(run, cn_font="黑体", size=18, bold=True)

    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=24)
    run = p.add_run('[目录请在 Word 中通过"引用"→"目录"→"自动目录"功能生成]')
    set_font(run, cn_font="宋体", size=12, color=(128, 128, 128))
    add_page_break(doc)

    # 8. 正文章节框架
    chapters = info.get("chapters", [
        {"num": 1, "title": "绪论", "sections": [
            {"num": "1.1", "title": "研究背景及意义"},
            {"num": "1.2", "title": "国内外研究现状"},
            {"num": "1.3", "title": "研究方法与内容"},
        ]},
        {"num": 2, "title": "需求分析", "sections": [
            {"num": "2.1", "title": "功能需求"},
            {"num": "2.2", "title": "非功能需求"},
            {"num": "2.3", "title": "可行性分析"},
        ]},
        {"num": 3, "title": "系统设计", "sections": [
            {"num": "3.1", "title": "系统总体结构设计"},
            {"num": "3.2", "title": "数据库设计"},
        ]},
        {"num": 4, "title": "详细设计与实现", "sections": [
            {"num": "4.1", "title": "各功能模块实现"},
        ]},
        {"num": 5, "title": "系统测试", "sections": [
            {"num": "5.1", "title": "测试概述"},
            {"num": "5.2", "title": "模块功能测试"},
            {"num": "5.3", "title": "测试结论"},
        ]},
    ])

    for ch in chapters:
        add_chapter_heading(doc, ch["num"], ch["title"])
        for sec in ch.get("sections", []):
            add_section_heading(doc, sec["num"], sec["title"])
            add_body_text(doc, f"[在此填写 {sec['num']} {sec['title']} 的内容]")
        add_page_break(doc)

    # 总结与展望
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=12)
    run = p.add_run("总结与展望")
    set_font(run, cn_font="黑体", size=18, bold=True)
    add_body_text(doc, "[在此填写总结与展望的内容]")
    add_page_break(doc)

    # 参考文献
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=12)
    run = p.add_run("参考文献")
    set_font(run, cn_font="黑体", size=18, bold=True)

    refs = info.get("references", [
        "[1] [在此填写参考文献，格式遵循 GB/T 7714]",
    ])
    for ref in refs:
        p = doc.add_paragraph()
        set_paragraph_format(p, line_spacing=1.5)
        run = p.add_run(ref)
        set_font(run, cn_font="宋体", size=10.5)
    add_page_break(doc)

    # 致谢
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=12)
    run = p.add_run("致 谢")
    set_font(run, cn_font="黑体", size=18, bold=True)
    add_body_text(doc, "[在此填写致谢内容]")

    return doc


def main():
    if len(sys.argv) < 3:
        print("用法: python generate_docx.py <info.json> <output.docx>")
        print("\ninfo.json 格式:")
        print(json.dumps({
            "title": "论文中文题目",
            "title_en": "English Title",
            "name": "姓名",
            "college": "学院",
            "major": "专业",
            "advisor": "指导教师",
            "date": "2025 年 5 月 25 日",
            "abstract_cn": "中文摘要",
            "abstract_en": "English abstract",
            "keywords_cn": "关键词1；关键词2；关键词3",
            "keywords_en": "keyword1; keyword2; keyword3",
            "chapters": [
                {"num": 1, "title": "绪论", "sections": [
                    {"num": "1.1", "title": "研究背景"}
                ]}
            ],
            "references": ["[1] 参考文献"]
        }, ensure_ascii=False, indent=2))
        sys.exit(1)

    info_file = sys.argv[1]
    output_file = sys.argv[2]

    with open(info_file, 'r', encoding='utf-8') as f:
        info = json.load(f)

    doc = generate_thesis(info)
    doc.save(output_file)
    print(f"论文模板已生成: {output_file}")


if __name__ == "__main__":
    main()
